#!/usr/bin/env python3
"""Estrae testo da file .docx"""

import sys
import io
from pathlib import Path
from docx import Document

def read_docx(file_path):
    """Legge un file .docx e stampa il testo"""
    try:
        sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
        doc = Document(file_path)
        for para in doc.paragraphs:
            if para.text.strip():
                print(para.text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        print(cell.text)
    except Exception as e:
        print(f"Errore: {e}")

if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Uso: python read_docx.py <path-to-docx>")
        sys.exit(1)

    read_docx(sys.argv[1])
