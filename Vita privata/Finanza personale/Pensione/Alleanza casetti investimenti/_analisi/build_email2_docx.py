# -*- coding: utf-8 -*-
import re
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX

INK = RGBColor(0x1C, 0x24, 0x30)
SOFT = RGBColor(0x5A, 0x62, 0x70)
ACCENT = RGBColor(0x9A, 0x6B, 0x2C)

doc = Document()

normal = doc.styles['Normal']
normal.font.name = 'Calibri'
normal.font.size = Pt(11)
normal.font.color.rgb = INK

for i in range(1, 3):
    h = doc.styles[f'Heading {i}']
    h.font.name = 'Calibri'
    h.font.color.rgb = INK
    h.font.bold = True

doc.styles['Heading 1'].font.size = Pt(16)
doc.styles['Heading 2'].font.size = Pt(12)


def add_rich(paragraph, text, italic=False):
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if not part:
            continue
        run = paragraph.add_run(part[2:-2] if part.startswith('**') else part)
        run.italic = italic


def p(text, space_after=8, bullet=False):
    para = doc.add_paragraph(style='List Bullet' if bullet else None)
    para.paragraph_format.space_after = Pt(space_after)
    add_rich(para, text)
    return para


def h2(text):
    doc.add_heading(text, level=2)


note = doc.add_paragraph()
note.add_run(
    "Nota: bozza pronta da incollare come testo semplice in Gmail. Conferma prima l'invio."
).italic = True

doc.add_paragraph()
p("Oggetto: Grazie per i chiarimenti — riepilogo e richiesta di incontro a Châtillon (Valle d'Aosta)", space_after=12)

p("Gentile Sig.ra Casetti,")
p("La ringrazio per le risposte fornite. Riepilogo quanto ho capito, per essere sicuro di procedere sulla base corretta, e Le indico i punti che restano da chiarire.")

h2("Confermato dalla Sua risposta")
p("Il diritto a richiedere capitale o rendita sulla posizione PIP matura solo al raggiungimento dei 67 anni: il fatto che il montante superi €100.000 in un anno precedente (es. 2031, ai ritmi di versamento attuali) non fa scattare nulla in quel momento — conta solo il valore a 67 anni. Questo esclude quindi qualunque ipotesi di prelievo pieno anticipato prima di quella data: prima dei 67 anni posso accedere solo alle anticipazioni parziali (30% senza motivo, 75% prima casa), non all'intera posizione.", bullet=True)
p("La rendita, una volta erogata, viene rivalutata nel tempo.", bullet=True)
p("Il tetto di capitale liquidabile potrebbe essere salito al 60% (era 50% nelle mie ipotesi) per una possibile novità della legge di bilancio — Le chiederei conferma di quale percentuale si applica esattamente al mio caso.", bullet=True)

h2("Sulla base di questo, la mia intenzione")
p("Fermare i versamenti volontari sul PIP nel corso del 2026, per lasciare margine di sicurezza rispetto alla soglia dei €100.000 a 67 anni (una proiezione fatta sui Suoi stessi dati storici mostra che aspettare oltre il 2027-2028 renderebbe comunque il superamento della soglia inevitabile, anche fermandosi). Le chiederei conferma che questa lettura sia corretta.")

h2("Le domande che restano senza una risposta numerica")
p("La rendita netta mensile reale a 67 anni, in euro.", bullet=True)
p("L'importo netto che percepirei io e quello che percepirebbe mia moglie Megumi Nishiya, scegliendo la reversibilità.", bullet=True)
p("Il coefficiente di conversione corretto da applicare al mio caso (4,19% implicito nella simulazione, contro il 5,772% riportato altrove nello stesso documento).", bullet=True)
p("Il capitale minimo garantito a 67 anni sul PIP, se il valore di €152.856,91 è solo una proiezione.", bullet=True)
p("Perché il capitale minimo garantito a scadenza della polizza D'Oro (€62.077,86) è inferiore alla somma dei premi complessivi attesi (€62.500).", bullet=True)
p("La scomposizione della perdita di €3.900 in caso di riscatto anticipato di D'Oro (tra costi di caricamento, quota di copertura assicurativa e penale di uscita anticipata), e l'anno o il valore accumulato in cui il riscatto tornerebbe in pareggio con i premi versati.", bullet=True)

h2("Richiesta di incontro di persona")
p("Vista la quantità di calcoli ancora da chiudere, credo sia più efficiente un incontro di persona. Vivo a Châtillon, in Valle d'Aosta: Le chiederei di trasferire le mie posizioni all'agenzia più vicina al mio domicilio, come mi ha gentilmente proposto, e di fissare un appuntamento.")

p("La ringrazio ancora per la disponibilità e resto in attesa di un Suo cortese riscontro.")
p("Best,")
p("Giampaolo Padula")

doc.save("Mail_Casetti_seconda_risposta.docx")
print("saved")
