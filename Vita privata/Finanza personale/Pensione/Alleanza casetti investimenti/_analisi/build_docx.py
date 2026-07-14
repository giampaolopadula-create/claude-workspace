# -*- coding: utf-8 -*-
import re
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

INK = RGBColor(0x1C, 0x24, 0x30)
SOFT = RGBColor(0x5A, 0x62, 0x70)
ACCENT = RGBColor(0x9A, 0x6B, 0x2C)
POS = RGBColor(0x2F, 0x6E, 0x52)
NEG = RGBColor(0x8A, 0x3B, 0x32)

SHADE_WARN = "F2E6E3"
SHADE_INFO = "F1E4CD"
SHADE_GOOD = "E6EDE8"
SHADE_HEAD = "E8E4D6"
SHADE_ALT = "F5F3EC"

doc = Document()

# base font
normal = doc.styles['Normal']
normal.font.name = 'Calibri'
normal.font.size = Pt(10.5)
normal.font.color.rgb = INK

for i in range(1, 4):
    h = doc.styles[f'Heading {i}']
    h.font.name = 'Georgia'
    h.font.color.rgb = INK
    h.font.bold = True

doc.styles['Heading 1'].font.size = Pt(20)
doc.styles['Heading 2'].font.size = Pt(14)
doc.styles['Heading 2'].font.color.rgb = ACCENT
doc.styles['Heading 3'].font.size = Pt(11.5)


def set_cell_shading(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def add_rich(paragraph, text, base_color=None, size=None, italic=False):
    """Parse **bold** markers into runs."""
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if not part:
            continue
        bold = part.startswith('**') and part.endswith('**')
        content = part[2:-2] if bold else part
        run = paragraph.add_run(content)
        run.bold = bold
        run.italic = italic
        if base_color:
            run.font.color.rgb = base_color
        if size:
            run.font.size = Pt(size)


def p(text, color=None, size=None, italic=False, space_after=6, align=None):
    para = doc.add_paragraph()
    para.paragraph_format.space_after = Pt(space_after)
    if align:
        para.alignment = align
    add_rich(para, text, base_color=color, size=size, italic=italic)
    return para


def h1(text):
    doc.add_heading(text, level=1)


def h2(num, text):
    para = doc.add_paragraph()
    para.style = doc.styles['Heading 2']
    r1 = para.add_run(f"{num}  ")
    r1.font.color.rgb = ACCENT
    r2 = para.add_run(text)
    r2.font.color.rgb = INK
    return para


def h3(text):
    doc.add_heading(text, level=3)


def lede(text):
    p(text, color=SOFT, italic=True, size=10, space_after=10)


def callout(label, lines, shade=SHADE_INFO, label_color=ACCENT):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.cell(0, 0)
    set_cell_shading(cell, shade)
    cell.paragraphs[0].paragraph_format.space_after = Pt(4)
    lab = cell.paragraphs[0]
    lab_run = lab.add_run(label.upper())
    lab_run.bold = True
    lab_run.font.size = Pt(9)
    lab_run.font.color.rgb = label_color
    for i, line in enumerate(lines):
        para = cell.add_paragraph() if i > 0 or True else cell.paragraphs[0]
        para.paragraph_format.space_after = Pt(4)
        add_rich(para, line)
    for cellpara in cell.paragraphs:
        cellpara.paragraph_format.space_before = Pt(2)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return tbl


def data_table(caption_text, headers, rows, highlight=None):
    """highlight: dict {(row_idx, col_idx): 'pos'/'neg'/'accent'}"""
    cap = doc.add_paragraph()
    cap_run = cap.add_run(caption_text.upper())
    cap_run.bold = True
    cap_run.font.size = Pt(8.5)
    cap_run.font.color.rgb = SOFT
    cap.paragraph_format.space_after = Pt(2)

    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = tbl.rows[0].cells
    for i, htext in enumerate(headers):
        set_cell_shading(hdr_cells[i], SHADE_HEAD)
        para = hdr_cells[i].paragraphs[0]
        run = para.add_run(htext)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = SOFT
        if i > 0:
            para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    for r_idx, row in enumerate(rows):
        cells = tbl.add_row().cells
        if r_idx % 2 == 1:
            for c in cells:
                set_cell_shading(c, SHADE_ALT)
        for c_idx, val in enumerate(row):
            para = cells[c_idx].paragraphs[0]
            run = para.add_run(str(val))
            run.font.size = Pt(9.5)
            if c_idx > 0:
                para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            if highlight and (r_idx, c_idx) in highlight:
                kind = highlight[(r_idx, c_idx)]
                run.font.color.rgb = {'pos': POS, 'neg': NEG, 'accent': ACCENT}[kind]
                run.bold = True
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return tbl


def small(text):
    p(text, color=SOFT, size=8.5, space_after=10)


def glossary_table(entries):
    tbl = doc.add_table(rows=0, cols=2)
    tbl.style = 'Table Grid'
    tbl.autofit = False
    tbl.columns[0].width = Cm(3.6)
    tbl.columns[1].width = Cm(12.5)
    for term, definition in entries:
        row = tbl.add_row().cells
        row[0].width = Cm(3.6)
        row[1].width = Cm(12.5)
        run = row[0].paragraphs[0].add_run(term)
        run.bold = True
        run.font.color.rgb = ACCENT
        run.font.size = Pt(9.5)
        run2 = row[1].paragraphs[0].add_run(definition)
        run2.font.size = Pt(9.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(6)


# ============================================================ TITLE
h1("Alleata Previdenza & D'Oro 2014")
p("Mantenere o riscattare? — versione sintetica", color=SOFT, size=12.5, space_after=2)
p("Preparato per Giampaolo Padula (n. 25/05/1974) — versione aggiornata 10 luglio 2026 — "
  "orizzonte analizzato: 67 anni / 2041 — tasso conto deposito usato: 2,2% netto",
  color=SOFT, size=9, space_after=16)

# ============================================================ GLOSSARY
h2("Glossario", "Termini tecnici — consulta al bisogno")
glossary_table([
    ("PIP", "Il nome tecnico del tuo fondo pensione (“Alleata Previdenza” è il nome commerciale)."),
    ("COVIP", "L'ente pubblico che vigila sui fondi pensione italiani."),
    ("Posizione individuale", "Il valore accumulato oggi nel fondo pensione."),
    ("Rendita vitalizia", "Pagamento mensile a vita. Si interrompe alla morte: nulla agli eredi, salvo reversibilità."),
    ("Rendita reversibile", "Come sopra, ma continua (importo ridotto) al coniuge dopo la morte."),
    ("Riscatto (fondo pensione)", "Uscita anticipata da un fondo pensione — per legge, solo invalidità o disoccupazione oltre 12 mesi."),
    ("Anticipazione", "Prelievo parziale restando iscritti: fino al 30% “per qualsiasi motivo” dopo 8 anni."),
    ("Riscatto (D'Oro)", "Per una polizza vita, chiudere in anticipo e incassare — sempre possibile, con eventuale penale."),
    ("Deduzione fiscale", "Riduce il reddito imponibile. Nel tuo caso, aliquota 43% su €5.000/anno versati ≈ €2.150/anno recuperati."),
    ("Conto deposito", "Conto bancario che blocca una somma per un periodo in cambio di un interesse."),
])

doc.add_page_break()

# ============================================================ §1 — DATI
h2("§1", "I due prodotti — dati chiave")

data_table("Alleata Previdenza (PIP) vs D'Oro 2014",
    ["Voce", "Alleata Previdenza (PIP)", "D'Oro 2014"],
    [
        ["Versamento annuo", "€5.000", "€2.500"],
        ["Valore posizione (2025)", "€68.204", "riscatto €21.142"],
        ["Deduzione fiscale", "Sì, 43% ≈ €2.150/anno", "No"],
        ["Uscita anticipata (prima dei 67)", "No, salvo casi di legge", "Sì, sempre possibile"],
        ["A scadenza (67 / 2040)", "Rendita obbligatoria per ≥50%", "Capitale libero al 100%"],
    ],
)

p("**La regola chiave del PIP**: a 67 anni puoi prendere il 100% in capitale solo se la posizione resta "
  "sotto ~€100.000 (soglia stimata da Alleanza). Continuando a versare €5.000/anno, la superi già nel "
  "2031 — ma il diritto di chiedere qualsiasi liquidazione scatta **solo a 67 anni**, non quando la "
  "posizione tocca €100.000 lungo il percorso. Per restare sotto soglia a 67 anni serve fermarsi con i "
  "versamenti nel **2026-2027**.")

# ============================================================ §2 — CONFRONTO
h2("§2", "Le tre opzioni a confronto")
lede("Rendita vitalizia completa, contro le due strade per avere capitale libero. In tutti e tre gli "
     "scenari D'Oro resta acceso fino alla sua scadenza naturale (2040, €62.000 garantiti) — non conviene "
     "mai riscattarlo prima, perché quel valore include anche i premi ancora da versare, non solo la "
     "crescita.")

h3("Come funziona ciascuna opzione")
p("**Rendita vitalizia** — continui a versare €5.000/anno nel PIP fino a 67 anni (posizione €152.857), "
  "converti tutto in rendita insieme a D'Oro: €648/mese garantiti a vita, nessun capitale libero.")
p("**Percorso legale** — fermi i versamenti PIP nel 2027 (resti sotto la soglia dei €100.000, §1); il PIP "
  "resta bloccato al suo interno fino a 67 anni, quando lo ritiri per intero in capitale. Nel frattempo "
  "continui a pagare D'Oro (€2.500/anno) fino al 2040, e risparmi altri €2.500/anno per conto tuo in un "
  "conto deposito. Non richiede alcuna conferma da Alleanza: è sempre legale.")
p("**Uscita 2030 (ipotetica)** — continui a versare nel PIP fino al 2030 (posizione €97.646), poi ne "
  "ritiri l'intero importo — **ipotesi non confermata**, dipende dalla risposta di Alleanza sul prelievo "
  "anticipato. D'Oro resta comunque acceso fino al 2040 come nel percorso legale. Dal 2031 investi "
  "€2.500/anno in più (oltre al premio D'Oro) in un conto deposito.")

data_table("Confronto sintetico",
    ["Criterio", "Rendita vitalizia\n(continui fino a 67)", "Percorso legale\n(fermi PIP 2027)", "Uscita 2030\n(ipotetica)"],
    [
        ["Capitale / reddito a 67 anni", "€648/mese a vita", "€199.400 capitale", "€216.800 capitale"],
        ["A €800/mese dura fino a", "n/a — fissa a €648/mese, non regolabile", "~95 anni", "~99 anni"],
        ["Allo stesso importo, €648/mese, dura fino a", "sempre", "~105 anni", "~111 anni"],
        ["Deduzione fiscale fino a", "67 anni", "2027", "2030"],
        ["Trasferibile agli eredi", "No (salvo reversibilità)", "Sì, per intero", "Sì, per intero"],
        ["Richiede conferma Alleanza", "No", "No", "Sì — non confermato"],
        ["Costo annuo per te", "€5.000/anno fino a 67", "€7.500 fino al 2027 (di cui ~€2.150/anno restituiti da 730), poi €5.000 fino al 2040", "€7.500 fino al 2030 (di cui ~€2.150/anno restituiti da 730), poi €5.000 fino al 2040"],
    ],
)

h3("Esploso: da cosa sono composti i due totali")

data_table("Percorso legale — €199.400",
    ["Componente", "Valore a 67 anni", "% del totale"],
    [
        ["PIP fermo dal 2027 (bloccato al suo rendimento fino a 67)", "€96.900", "49%"],
        ["D'Oro tenuto a scadenza 2040 (capitale minimo garantito)", "€62.000", "31%"],
        ["€2.500/anno risparmiati in conto deposito dal 2028 al 2041 (14 anni, al 2,2%)", "€40.500", "20%"],
        ["Totale a 67 anni", "€199.400", "100%"],
    ],
    highlight={(3,1):'accent'}
)

data_table("Uscita 2030 (ipotetica) — €216.800",
    ["Componente", "Valore a 67 anni", "% del totale"],
    [
        ["PIP ritirato nel 2030 (€97.646), investito 11 anni al 2,2%", "€124.100", "57%"],
        ["D'Oro tenuto a scadenza 2040 (capitale minimo garantito, invariato)", "€62.000", "29%"],
        ["€2.500/anno investiti dal 2031 al 2041 (11 anni, al 2,2%)", "€30.700", "14%"],
        ["Totale a 67 anni", "€216.800", "100%"],
    ],
    highlight={(3,1):'accent'}
)

callout("Tenere D'Oro fino al 2040 è una scelta di sicurezza, non di rendimento massimo", [
    "A parità di soldi versati, riscattare D'Oro nel 2030 (~€35.000) e investire anche i €2.500/anno "
    "successivi in conto deposito darebbe circa **€75.200** entro il 2041 — più dei €62.000 del capitale "
    "minimo garantito a scadenza. Il “minimo garantito” di D'Oro, infatti, è quasi un pareggio puro con la "
    "somma dei premi totali versati: da solo non batte nemmeno l'inflazione.",
    "**Perché nell'esploso tengo comunque D'Oro fino al 2040**: è la scelta più sicura e tranquilla, non "
    "quella che massimizza il rendimento atteso. D'Oro garantisce un pavimento contro rendimenti negativi "
    "e paga il capitale ai beneficiari fuori successione e rapidamente in caso di decesso — una protezione "
    "familiare che un conto deposito non replica allo stesso modo. Rinunciare a ~€13.000 di rendimento "
    "atteso per questa tranquillità è una scelta legittima, coerente con l'approccio prudente scelto per "
    "tutto il resto del piano."
], shade=SHADE_GOOD, label_color=POS)

# ============================================================ §3 — PAREGGIO E REVERSIBILITÀ
h2("§3", "Punto di pareggio e costo della mancata reversibilità")
lede("Se converti l'intera posizione PIP (€152.857) in rendita vitalizia a 67 anni (€6.401,93 lordi/anno).")

data_table("Quanto “vale” la rendita a seconda di quanto vivi",
    ["Se muori a", "Rendita incassata (lorda)", "% del capitale convertito", "Capitale perso per gli eredi"],
    [
        ["75 anni", "€51.215", "33%", "€101.642"],
        ["85 anni", "€115.235", "75%", "€37.622"],
        ["91 anni", "€153.646", "101% — pareggio", "≈ €0"],
        ["95 anni", "€179.254", "117%", "in vantaggio"],
    ],
)

p("Il pareggio è a **~91 anni**. Prima di quell'età, ogni anno di vita in meno è capitale che resta alla "
  "compagnia. La rendita **reversibile** (a favore di Megumi) risolverebbe il problema ma riduce l'importo "
  "mensile — Alleanza non ha ancora dato il coefficiente esatto.")

# ============================================================ SINTESI FINALE
doc.add_page_break()
h2("Sintesi", "Cosa fare, in ordine")

callout("1 — Nessuna azione richiesta prima del 2027", [
    "Continuare come oggi (€5.000 PIP + €2.500 D'Oro) è coerente con tutte le strade possibili almeno fino "
    "al 2026-2027."
], shade=SHADE_INFO, label_color=ACCENT)

callout("2 — Servono due risposte scritte da Casetti prima di decidere", [
    "**a)** Il prelievo del 100% della posizione PIP sotto soglia è possibile solo a 67 anni, o anche prima "
    "se la posizione resta sotto €100.000? (Se sì → Uscita 2030, €216.800. Se no → Percorso legale, €199.400.)",
    "**b)** Una volta iniziata l'erogazione, la rendita resta fissa o si rivaluta nel tempo?",
    "**c)** Qual è il coefficiente della rendita reversibile a favore di Megumi?"
], shade=SHADE_WARN, label_color=NEG)

callout("3 — In entrambi gli scenari di capitale libero, la rendita vitalizia conviene solo oltre i 95-99 anni", [
    "Sotto quell'età, sia il Percorso legale sia l'Uscita 2030 generano più reddito mensile e restano "
    "trasferibili agli eredi. Oltre quell'età, solo la rendita continua a pagare."
], shade=SHADE_GOOD, label_color=POS)

callout("4 — La deduzione fiscale resta l'argomento più forte per non fermarsi subito", [
    "Il 43% su €5.000/anno è un guadagno certo e immediato che nessun conto deposito replica — vale la pena "
    "continuare a versare almeno fino al 2026-2027, indipendentemente dalla strada scelta poi."
], shade=SHADE_INFO, label_color=ACCENT)

# ============================================================ APPENDIX: email suggestion
doc.add_page_break()
h2("Appendice", "Suggerimento per la mail a Casetti")
lede("Nella bozza “Ultima Mail di risposta alla sig.ra Casetti”, punto A, la domanda sull'anno di "
     "superamento soglia rischia una risposta ambigua. Frase da aggiungere subito dopo:")

callout("Frase suggerita", [
    "“Vorrei essere sicuro di aver capito correttamente il meccanismo: il superamento della soglia dei "
    "€100.000 in un dato anno (es. 2031) non mi dà di per sé alcun diritto di richiedere la liquidazione in "
    "quel momento, corretto? Il diritto di richiedere la prestazione (in capitale o rendita) nasce solo al "
    "raggiungimento dell'età pensionabile (67 anni), e solo in quel momento si guarda se la posizione è sopra "
    "o sotto soglia. È corretto? Se sì, l'unica leva che ho oggi è decidere se interrompere i versamenti ora, "
    "in modo che la posizione, dopo la sola rivalutazione fino al 2041, resti sotto soglia in quell'anno — "
    "non aspettare che tocchi €100.000 lungo il percorso.”"
], shade=SHADE_INFO, label_color=ACCENT)

# ============================================================ FOOTER
foot = doc.add_paragraph()
foot.paragraph_format.space_before = Pt(16)
run = foot.add_run(
    "Fonti: prospetti Alleata Previdenza (2023/2024/2025), estratti conto e simulazione personalizzata "
    "(01/07/2026), estratto conto D'Oro 2014 (04/03/2025), corrispondenza con Claudia Casetti e Michelina "
    "Tenore (Alleanza Assicurazioni, Agenzia Roma San Giovanni). I calcoli indicati come stime sono "
    "elaborazioni proprie non ufficiali; da verificare con una simulazione ufficiale Alleanza prima di "
    "qualunque decisione operativa. La rendita D'Oro è una stima illustrativa: Alleanza ha dichiarato che "
    "il calcolo reale sarà possibile solo alla scadenza (2040)."
)
run.font.size = Pt(8)
run.font.color.rgb = SOFT
run.italic = True

doc.save("Alleata_Previdenza_DOro_report.docx")
print("saved")
