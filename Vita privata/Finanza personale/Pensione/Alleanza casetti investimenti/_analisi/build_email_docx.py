# -*- coding: utf-8 -*-
import re
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX

INK = RGBColor(0x1C, 0x24, 0x30)
SOFT = RGBColor(0x5A, 0x62, 0x70)
ACCENT = RGBColor(0x9A, 0x6B, 0x2C)

doc = Document()

normal = doc.styles['Normal']
normal.font.name = 'Calibri'
normal.font.size = Pt(11)
normal.font.color.rgb = INK

for i in range(1, 3):
    h = doc.styles[f'Heading {i}']
    h.font.name = 'Calibri'
    h.font.color.rgb = INK
    h.font.bold = True

doc.styles['Heading 1'].font.size = Pt(16)
doc.styles['Heading 2'].font.size = Pt(12)


def add_rich(paragraph, text, italic=False):
    """Parse **bold** markers into runs; bold runs rendered highlighted (new text)."""
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if not part:
            continue
        is_new = part.startswith('**') and part.endswith('**')
        content = part[2:-2] if is_new else part
        run = paragraph.add_run(content)
        run.italic = italic
        if is_new:
            run.bold = True
            run.font.highlight_color = WD_COLOR_INDEX.YELLOW


def p(text, size=11, italic=False, space_after=8, color=None, align=None):
    para = doc.add_paragraph()
    para.paragraph_format.space_after = Pt(space_after)
    if align:
        para.alignment = align
    add_rich(para, text, italic=italic)
    if color:
        for r in para.runs:
            if not r.font.highlight_color:
                r.font.color.rgb = color
    return para


def bullet(text):
    para = doc.add_paragraph(style='List Bullet')
    para.paragraph_format.space_after = Pt(6)
    add_rich(para, text)
    return para


def bold_label(text):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(10)
    para.paragraph_format.space_after = Pt(4)
    run = para.add_run(text)
    run.bold = True
    return para


# ============================================================ NOTE
note = doc.add_paragraph()
note.paragraph_format.space_after = Pt(14)
r = note.add_run(
    "Nota: il testo evidenziato in giallo è l'unica parte nuova aggiunta rispetto alla tua bozza originale. "
    "Il resto è identico. Puoi rimuovere l'evidenziazione (o lasciarla, non compare nell'email inviata se la "
    "incolli come testo semplice) prima di inviare."
)
r.italic = True
r.font.size = Pt(9)
r.font.color.rgb = SOFT

doc.add_heading("Oggetto: Chiarimenti numerici su soglia €100.000, coefficienti, rendita netta e polizza D'Oro", level=2)

p("Gentile Sig.ra Casetti,")

p("La ringrazio per le risposte fornite finora. Sono nato il 25 maggio 1974 e ho quindi oggi 52 anni: ho "
  "provato a fare io stesso una proiezione della mia posizione sulla base dei dati storici disponibili, per "
  "capire meglio i tempi reali legati alla soglia dei €100.000, e Le chiederei di confermarla o correggerla "
  "con i Suoi strumenti, oltre a darmi alcuni ultimi chiarimenti operativi.")

bold_label("A) Soglia dei €100.000 e strategia di versamento")

p("Dal prospetto annuale risulta che la mia posizione è cresciuta così negli ultimi anni:")
bullet("31/12/2023: €56.855,21 (versato cumulato €55.006,00)")
bullet("31/12/2024: €62.477,00 (incremento nell'anno: €5.621,79)")
bullet("31/12/2025: €68.204,07 (incremento nell'anno: €5.727,07)")

p("Prendendo come riferimento un contributo netto annuo di circa €4.771 (al netto dei costi, in linea con gli "
  "ultimi tre anni) e un rendimento netto medio dell'1,4% (in linea con quello certificato negli ultimi "
  "esercizi, tra l'1,13% e l'1,46%), la mia proiezione anno per anno è la seguente:")

tbl = doc.add_table(rows=1, cols=3)
tbl.style = 'Table Grid'
hdr = tbl.rows[0].cells
for i, t in enumerate(["Anno", "Posizione stimata", "Mia età (31/12)"]):
    run = hdr[i].paragraphs[0].add_run(t)
    run.bold = True
    run.font.size = Pt(10)
rows = [
    ("2026", "€73.929,93", "52"),
    ("2027", "€79.735,95", "53"),
    ("2028", "€85.623,25", "54"),
    ("2029", "€91.592,97", "55"),
    ("2030", "€97.646,28", "56"),
    ("2031", "€103.784,32", "57"),
]
for row in rows:
    cells = tbl.add_row().cells
    for i, v in enumerate(row):
        run = cells[i].paragraphs[0].add_run(v)
        run.font.size = Pt(10)
        if i > 0:
            cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
doc.add_paragraph().paragraph_format.space_after = Pt(6)

p("Se questi numeri sono corretti, supererei la soglia dei €100.000 già durante il 2031, a 56-57 anni: circa "
  "dieci anni prima del pensionamento previsto a 67 anni (2041), non “in prossimità dell'età "
  "pensionabile” come lascia intendere la Sua risposta precedente.")

p("Le chiederei quindi:")

bullet("Può confermare o correggere questa proiezione con il calcolo ufficiale di Alleanza, indicando l'anno "
       "o l'età esatta in cui, ai ritmi di versamento attuali, supererei la soglia utile per il 100% capitale?")

bullet("**Vorrei essere sicuro di aver capito correttamente il meccanismo: il superamento della soglia dei "
       "€100.000 in un dato anno (es. 2031) non mi dà di per sé alcun diritto di richiedere la liquidazione "
       "in quel momento, corretto? Il diritto di richiedere la prestazione (in capitale o rendita) nasce solo "
       "al raggiungimento dell'età pensionabile (67 anni), e solo in quel momento si guarda se la posizione è "
       "sopra o sotto soglia. È corretto? Se sì, l'unica leva che ho oggi è decidere se interrompere i "
       "versamenti ora, in modo che la posizione, dopo la sola rivalutazione fino al 2041, resti sotto soglia "
       "in quell'anno — non aspettare che tocchi €100.000 lungo il percorso.**")

bullet("Potrebbe indicarmi il calcolo analitico dietro la soglia di “circa €100.000”: assegno "
       "sociale INPS di riferimento, il 50% di tale importo, il coefficiente di conversione applicato e "
       "l'età presa a riferimento? Vorrei capire se la soglia stessa si sposta nel tempo (per rivalutazione "
       "dell'assegno sociale o aggiornamento dei coefficienti) o se resta stabile.")

bullet("Una volta superata la soglia, cosa succede esattamente ai versamenti successivi: continuano comunque "
       "ad aumentare la posizione (con la parte eccedente obbligatoriamente destinata a rendita), oppure "
       "conviene interrompere i versamenti proprio in prossimità di quel momento?")

bullet("Le chiederei una simulazione per tre scenari concreti: (a) fermarmi al momento in cui la posizione "
       "raggiunge la soglia utile per il 100% capitale; (b) proseguire fino a 67 anni; (c) proseguire fino a "
       "68-71 anni. Per ciascuno: quota liquidabile subito in capitale e quota da convertire in rendita.")

bullet("Dato che a 67 anni la mia posizione risulterà comunque ben oltre la soglia dei €100.000 (proiezione: "
       "€152.856,91), la quota liquidabile subito in capitale resta fissa (es. 50%) o si riduce ulteriormente "
       "proseguendo i versamenti fino a 68-71 anni? In sostanza: superata la soglia, continuare a versare "
       "porta ancora un vantaggio concreto oltre alla deduzione fiscale, o converrebbe destinare i versamenti "
       "successivi ad altri strumenti?")

bold_label("B) Coefficiente di conversione e garanzia")

p("Nella simulazione a 67 anni, €152.856,91 di posizione generano €6.401,93 di rendita, cioè un coefficiente "
  "implicito del 4,19%. Nello stesso documento compare però un esempio con coefficiente 5,772%. Quale "
  "coefficiente si applica realmente al mio caso, e perché tale differenza?")

p("Il valore di €152.856,91 è garantito contrattualmente o è solo una proiezione? Se è una proiezione, qual "
  "è invece il capitale minimo garantito a 67 anni?")

bold_label("C) Rendita netta e reversibilità")

p("Quale sarebbe l'importo netto mensile reale (in euro, non l'aliquota) della rendita a 67 anni?")

p("Scegliendo la reversibilità a favore di mia moglie Megumi Nishiya, quanto percepirei io netto al mese e "
  "quanto percepirebbe lei dopo di me?")

p("Qual è il punto di equilibrio della rendita: trasformando il capitale in rendita, per quanti anni dovrei "
  "percepirla per recuperare quel capitale?")

p("**Una volta iniziata l'erogazione, l'importo mensile della rendita resta fisso o viene rivalutato negli "
  "anni (ad esempio in base al rendimento della gestione separata)? Vorrei capire se, dopo 20-25 anni di "
  "rendita, l'importo percepito manterrà il potere d'acquisto di oggi o resterà fisso in euro nominali, "
  "perdendo progressivamente valore reale con l'inflazione.**")

bold_label("D) Rendimenti storici")

p("Potrebbe fornirmi lo storico dei rendimenti netti effettivamente riconosciuti, anno per anno, sia per "
  "Alleata Previdenza sia per la polizza D'Oro, dalla rispettiva decorrenza a oggi?")

bold_label("E) Polizza D'Oro")

p("Anche in forma indicativa e non vincolante (con i coefficienti oggi in vigore), potrebbe fornirmi una "
  "simulazione della rendita netta mensile a scadenza (2040), reversibile e non?")

p("Il 3,15% lordo del fondo Euro San Giorgio: quanto arriva effettivamente, netto, alla mia polizza, anno "
  "per anno?")

p("Perché il valore garantito a scadenza (€62.077,86) è inferiore alla somma dei premi complessivi attesi "
  "(€2.500 × 25 anni = €62.500)?")

p("Sulla perdita di €3.900 in caso di riscatto: potrebbe scomporla tra costi/caricamenti, quota copertura "
  "assicurativa e penalità di riscatto anticipato?")

p("La differenza attuale tra premi versati e valore di riscatto (-€3.900) da cosa dipende principalmente: "
  "costi di caricamento iniziali che si esauriscono nel tempo, o una componente che incide per tutta la "
  "durata del contratto? A che anno (o a quale valore accumulato) è previsto che il valore di riscatto torni "
  "a superare il totale dei premi versati, se questo è previsto prima della scadenza?")

bold_label("F) Beneficiari")

p("Le chiederei conferma scritta formale, per entrambe le posizioni, di nome, cognome, codice fiscale, "
  "percentuale e prodotto associato al beneficiario designato (Sig.ra Megumi Nishiya).")

p("La ringrazio per la disponibilità e resto in attesa di un Suo cortese riscontro.")

p("Cordiali saluti,")
p("Gp")

doc.save("Mail_Casetti_revisionata.docx")
print("saved")
